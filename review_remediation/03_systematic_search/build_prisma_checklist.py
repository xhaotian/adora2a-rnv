#!/usr/bin/env python3
"""Create a completed PRISMA 2020 checklist keyed to manuscript sections."""
from pathlib import Path
from docx import Document
from docx.shared import Inches

PROJECT=Path(__file__).resolve().parents[2];FINAL=PROJECT/"final_submission"
items=[
("1","Title","Identify the report as a systematic review.","Title; Abstract"),
("2","Abstract","PRISMA abstract elements.","Abstract"),
("3","Rationale","Describe the rationale in context.","Introduction"),
("4","Objectives","Provide explicit objectives.","Introduction, final paragraph"),
("5","Eligibility criteria","Specify inclusion and exclusion criteria.","Methods: Systematic dataset identification and screening"),
("6","Information sources","Specify databases, registers, websites and date last searched.","Methods; Supplementary Methods; Search log"),
("7","Search strategy","Present full strategies for all sources.","Supplementary Methods; Search_All_Sources.tsv"),
("8","Selection process","Describe how records were screened.","Methods: Systematic dataset identification and screening"),
("9","Data collection","Describe data collection and confirmation.","Methods: Mouse expression reconstruction; Source Data"),
("10a","Data items—outcomes","List and define outcomes.","Methods: Mouse meta-analysis"),
("10b","Data items—other","List other variables and assumptions.","Methods: Biological-unit definitions; Supplementary Table S1"),
("11","Risk of bias","Describe the method used to assess risk of bias in included studies.","Methods: Dataset-level evidence-quality appraisal; no clinical/intervention instrument was directly applicable; prespecified dataset-level domains were assessed"),
("12","Effect measures","Specify effect measure.","Methods: Mouse meta-analysis"),
("13a","Synthesis—eligibility","Describe processes deciding study eligibility for synthesis.","Methods: eligibility and duplicate-cohort audit"),
("13b","Synthesis—preparation","Describe data preparation.","Methods: Mouse expression reconstruction"),
("13c","Synthesis—display","Describe tabulation and visual display.","Methods; Fig 2"),
("13d","Synthesis—methods","Describe statistical synthesis.","Methods: Mouse meta-analysis"),
("13e","Heterogeneity","Describe exploration of heterogeneity.","Methods; leave-one-accession-out and prediction interval"),
("13f","Sensitivity","Describe sensitivity analyses.","Methods: strict biological-unit sensitivity"),
("14","Reporting bias","Describe assessment or reason not performed.","Methods: Mouse meta-analysis"),
("15","Certainty","Describe assessment of certainty.","Evidence boundaries in Results and Discussion; no formal GRADE"),
("16a","Study selection","Report search and selection results.","Results; Fig 1; PRISMA_Flow_Source.tsv"),
("16b","Excluded studies","Cite exclusions that might appear eligible.","Results; screening log"),
("17","Study characteristics","Report characteristics of included studies.","Table 1; Supplementary Table S1"),
("18","Risk of bias","Present assessments for each included study.","Supplementary Table S1: dataset-level evidence-quality table; Source Data; Discussion"),
("19","Individual results","Present effect and precision for each study.","Fig 2A; Source Data"),
("20a","Synthesis contributors","Summarize contributing studies.","Results: mouse synthesis"),
("20b","Statistical synthesis","Present pooled effect and heterogeneity.","Results; Fig 2"),
("20c","Heterogeneity investigations","Present leave-one-out results.","Fig 2B; Source Data"),
("20d","Sensitivity results","Present strict-unit sensitivity.","Results; Fig 2C"),
("21","Reporting biases","Present assessments of reporting bias.","Methods; Results; S6 Fig: descriptive funnel plot and exploratory Pustejovsky–Rodgers diagnostic"),
("22","Certainty","Present evidence limitations.","Discussion"),
("23a","Interpretation","Interpret results in context.","Discussion"),
("23b","Evidence limitations","Discuss limitations of included evidence.","Discussion"),
("23c","Review limitations","Discuss limitations of review processes.","Discussion"),
("23d","Implications","Discuss implications.","Discussion; Conclusions"),
("24a","Registration","Provide registration information.","No protocol was prospectively registered"),
("24b","Protocol","Indicate where protocol can be accessed.","Not applicable; no prospective protocol"),
("24c","Amendments","Describe amendments.","Not applicable"),
("25","Support","Describe financial or non-financial support.","Submission-system funding field"),
("26","Competing interests","Declare competing interests.","Submission-system metadata"),
("27","Availability","Report availability of data, code and materials.","Data availability; Code availability; Source Data"),
]
d=Document();d.add_heading("PRISMA 2020 checklist",0);d.add_paragraph("Review component: systematic identification and synthesis of P17 mouse oxygen-induced-retinopathy transcriptomic datasets. Search date: 11 August 2026.")
t=d.add_table(rows=1,cols=4);t.style="Table Grid";hdr=t.rows[0].cells
for i,x in enumerate(["Item","Section/topic","Checklist item","Location"]):hdr[i].text=x
for row in items:
 c=t.add_row().cells
 for i,x in enumerate(row):c[i].text=x
d.sections[0].left_margin=Inches(.55);d.sections[0].right_margin=Inches(.55)
FINAL.mkdir(parents=True,exist_ok=True);d.save(FINAL/"PRISMA_Checklist_FINAL.docx")
print(FINAL/"PRISMA_Checklist_FINAL.docx")
